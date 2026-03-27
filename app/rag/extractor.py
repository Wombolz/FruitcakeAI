"""
FruitcakeAI v5 — Document extraction utilities.

Owns file-type specific extraction before text reaches the indexing layer.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

try:
    import fitz  # PyMuPDF

    HAS_PYMUPDF = True
except ImportError:  # pragma: no cover - optional dependency
    fitz = None  # type: ignore[assignment]
    HAS_PYMUPDF = False

try:
    from pypdf import PdfReader

    HAS_PYPDF = True
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None  # type: ignore[assignment]
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None  # type: ignore[assignment]
    HAS_DOCX = False

try:
    import pytesseract
    from PIL import Image

    HAS_TESSERACT = True
except ImportError:  # pragma: no cover - optional dependency
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]
    HAS_TESSERACT = False


log = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Raised when a document cannot be extracted into usable text."""


class DocumentExtractor:
    _TEXT_EXTENSIONS = {
        ".txt": "txt",
        ".md": "md",
        ".py": "code",
        ".yaml": "config",
        ".yml": "config",
        ".json": "config",
        ".toml": "config",
        ".swift": "code",
        ".js": "code",
        ".ts": "code",
        ".jsx": "code",
        ".tsx": "code",
        ".html": "html",
        ".css": "code",
        ".sh": "code",
        ".sql": "code",
        ".xml": "config",
        ".ini": "config",
        ".env": "config",
    }
    _SPECIAL_TEXT_FILENAMES = {
        "dockerfile": "config",
        "makefile": "config",
        "justfile": "config",
        "agents.md": "md",
        "skill.md": "md",
    }

    def supports(self, file_path: Path) -> bool:
        try:
            self.content_type_from_extension(file_path)
            return True
        except ExtractionError:
            return False

    def content_type_from_extension(self, file_path: Path) -> str:
        special = self._SPECIAL_TEXT_FILENAMES.get(file_path.name.lower())
        if special:
            return special
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        if suffix == ".docx":
            return "docx"
        mapped = self._TEXT_EXTENSIONS.get(suffix)
        if mapped:
            return mapped
        raise ExtractionError(f"Unsupported file format: {suffix or '<none>'}")

    def extract(self, file_path: Path) -> tuple[str, str]:
        if not file_path.exists():
            raise ExtractionError(f"File not found: {file_path}")

        content_type = self.content_type_from_extension(file_path)
        if content_type == "pdf":
            return self._extract_pdf(file_path)
        if content_type == "docx":
            return self._extract_docx(file_path)
        if content_type == "md":
            return self._extract_markdown(file_path)
        if content_type in {"txt", "code", "config", "html"}:
            return self._extract_plaintext(file_path)
        raise ExtractionError(f"Unsupported content type: {content_type}")

    def _read_text(self, file_path: Path) -> str:
        for encoding in ("utf-8", "latin-1"):
            try:
                return file_path.read_text(encoding=encoding).strip()
            except UnicodeDecodeError:
                continue
            except Exception as exc:  # pragma: no cover - rare file errors
                raise ExtractionError(f"Failed reading text file: {exc}") from exc
        raise ExtractionError(f"Could not decode text file: {file_path.name}")

    def _extract_plaintext(self, file_path: Path) -> tuple[str, str]:
        text = self._read_text(file_path)
        if not text:
            raise ExtractionError("No text extracted from plaintext file")
        return "plaintext", text

    def _extract_markdown(self, file_path: Path) -> tuple[str, str]:
        text = self._read_text(file_path)
        if text.startswith("---"):
            lines = text.splitlines()
            if len(lines) > 1:
                for idx in range(1, len(lines)):
                    if lines[idx].strip() == "---":
                        text = "\n".join(lines[idx + 1 :]).strip()
                        break
        if not text:
            raise ExtractionError("No text extracted from markdown file")
        return "markdown", text

    def _extract_docx(self, file_path: Path) -> tuple[str, str]:
        if not HAS_DOCX or DocxDocument is None:
            raise ExtractionError("python-docx is not available")
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables.append(" | ".join(row_text))
            text = "\n\n".join(paragraphs)
            if tables:
                text = f"{text}\n\n--- Tables ---\n\n" + "\n".join(tables) if text else "\n".join(tables)
            text = text.strip()
            if not text:
                raise ExtractionError("No text extracted from DOCX file")
            return "python-docx", text
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError(f"Failed reading DOCX file: {exc}") from exc

    def _extract_pdf(self, file_path: Path) -> tuple[str, str]:
        if HAS_PYMUPDF:
            try:
                return "pymupdf", self._extract_pdf_pymupdf(file_path)
            except ExtractionError as exc:
                log.warning("PyMuPDF extraction failed; falling back to pypdf", file=str(file_path), error=str(exc))
        if HAS_PYPDF:
            try:
                return "pypdf2", self._extract_pdf_pypdf(file_path)
            except ExtractionError as exc:
                log.warning("pypdf extraction failed", file=str(file_path), error=str(exc))
        if HAS_TESSERACT and HAS_PYMUPDF:
            try:
                return "ocr", self._extract_pdf_ocr(file_path)
            except ExtractionError as exc:
                log.warning("OCR extraction failed", file=str(file_path), error=str(exc))
        raise ExtractionError("No usable PDF extraction backend available")

    def _extract_pdf_pymupdf(self, file_path: Path) -> str:
        try:
            doc = fitz.open(str(file_path))  # type: ignore[union-attr]
            text_parts: list[str] = []
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text("text").strip()
                if not text:
                    blocks = page.get_text("blocks")
                    text = "\n".join(
                        block[4].strip() for block in blocks if len(block) > 4 and str(block[4]).strip()
                    ).strip()
                if text:
                    text_parts.append(text)
            extracted = self._clean_pdf_text("\n\n".join(text_parts).strip())
            if not extracted:
                raise ExtractionError("PyMuPDF extracted no text")
            return extracted
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError(f"Failed reading PDF with PyMuPDF: {exc}") from exc

    def _extract_pdf_pypdf(self, file_path: Path) -> str:
        try:
            reader = PdfReader(str(file_path))  # type: ignore[operator]
            pages: list[str] = []
            for page in reader.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
            extracted = self._clean_pdf_text("\n\n".join(pages).strip())
            if not extracted:
                raise ExtractionError("pypdf extracted no text")
            return extracted
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError(f"Failed reading PDF with pypdf: {exc}") from exc

    def _extract_pdf_ocr(self, file_path: Path) -> str:
        try:
            doc = fitz.open(str(file_path))  # type: ignore[union-attr]
            mat = fitz.Matrix(2.0, 2.0)  # type: ignore[union-attr]
            parts: list[str] = []
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))  # type: ignore[union-attr]
                text = pytesseract.image_to_string(img).strip()  # type: ignore[union-attr]
                if text:
                    parts.append(f"[Page {page_num + 1}]\n{text}")
            extracted = "\n\n".join(parts).strip()
            if not extracted:
                raise ExtractionError("OCR extracted no text")
            return extracted
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError(f"Failed OCR extraction for PDF: {exc}") from exc

    def _clean_pdf_text(self, text: str) -> str:
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        lines = text.splitlines()
        cleaned: list[str] = []
        seen: set[str] = set()
        for line in lines:
            stripped = line.strip()
            if not stripped:
                cleaned.append("")
                continue
            if stripped not in seen or len(stripped) < 20:
                cleaned.append(stripped)
                seen.add(stripped)
        return "\n".join(cleaned).strip()
